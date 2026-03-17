import os
import time
import json
import re
import requests
import pdfplumber
import streamlit as st
from datetime import datetime
import config  # 引用你现有的配置文件
from doc_generator import DocGenerator
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH # 新增：用于控制两端对齐
import io
import re


# --- WSH(Wall Street Highlight) ---
STEP_1_PROMPT_TEMPLATE = config.STEP_1_PROMPT_TEMPLATE
STEP_2_PROMPT_TEMPLATE = config.STEP_2_PROMPT_TEMPLATE

# --- 资金流周报 (Weekly Fund Flow) ---
FUND_FLOW_STEP1 = config.FUND_FLOW_STEP1
FUND_FLOW_STEP2 = config.FUND_FLOW_STEP2



# 预设模型列表 (供 Tab 2 使用)
AVAILABLE_MODELS = [config.AI_MODEL_NAME,  "gemini-3-pro-preview","claude-sonnet-4","claude-sonnet-4-5", "gemini-2.5-flash","gemini-2.5-pro", "deepseek-v3.2-exp",
    "deepseek-r1",
    "qwen3-max",
    "qwen3-235b-a22b-thinking-2507",
    "qwen3-235b-a22b-instruct-2507",
    "qwen3-vl-plus",
    "qwen3-vl-flash",
    "kimi-k2-thinking",
    "Moonshot-Kimi-K2-Instruct",
    "glm-4.6"]


def get_bank_acronym(full_name):
    """
    根据 Step 1 提取的全名，返回对应的缩写
    """
    if not full_name: return "Unknown"
    name_upper = full_name.upper()
    if "J.P. MORGAN" in name_upper or "JPMORGAN" in name_upper: return "JPM"
    if "GOLDMAN" in name_upper: return "GS"
    if "MORGAN STANLEY" in name_upper: return "MS"
    if "DEUTSCHE" in name_upper: return "DB"
    if "CITIC" in name_upper: return "CITICS"
    if "AMERICA" in name_upper or "BOFA" in name_upper: return "BofA"
    if "UBS" in name_upper: return "UBS"
    if "HSBC" in name_upper: return "HSBC"

    clean_name = re.sub(r'[^\w]', '', full_name.split()[0])
    return clean_name

# ================= 功能函数 =================

def extract_pdf_text(path):
    print(f"📄 正在读取 PDF: {path}...")
    full_text = ""
    try:
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text: full_text += text + "\n"
        return full_text
    except Exception as e:
        print(f"❌ 读取 PDF 失败: {e}")
        return None

def get_token():
    payload = {'grant_type': 'client_credentials', 'client_id': config.CLIENT_ID, 'client_secret': config.CLIENT_SECRET}
    try:
        resp = requests.post(config.AUTH_URL, data=payload, timeout=10)
        resp.raise_for_status()
        return resp.json().get('access_token')
    except Exception as e:
        print(f"❌ Token 获取失败: {e}")
        return None

# 🌟 修改点：增加了 model_name 参数，以兼容 Tab 2 的自定义模型选择
def call_ai_and_wait_generic(system_prompt, user_content, model_name=None):
    token = get_token()
    if not token: return None

    full_prompt = f"{system_prompt}\n\n=== INPUT DATA ===\n{user_content}"
    url = f"{config.API_BASE_URL}/job"
    headers = {'Authorization': f'Bearer {token}', 'Content-Type': 'application/json'}

    actual_model = model_name if model_name else config.AI_MODEL_NAME

    payload = {
        "type": "callLlm",
        "metadata": config.API_METADATA,
        "input": {"parameter": {"model_name": actual_model, "prompt": full_prompt}}
    }

    try:
        print(f"🚀 提交 AI 任务 (Model: {actual_model})...")
        resp = requests.post(url, headers=headers, json=payload, timeout=30)
        if resp.status_code != 200:
            print(f"❌ 提交失败: {resp.text}")
            return None

        job_id = resp.json().get("id") or resp.json().get("uuid")
        print(f"⏳ 等待 AI (ID: {job_id})...")

        for i in range(60): 
            time.sleep(2)
            check_url = f"{config.API_BASE_URL}/job/JOB_ID/{job_id}"
            check_resp = requests.get(check_url, headers=headers)

            if check_resp.status_code == 200:
                res = check_resp.json()
                status = res.get("status")
                if status in ["SUCCESS", "COMPLETED"]:
                    print("✅ AI 完成！")
                    return clean_json(res.get("output") or res.get("result"))
                elif status == "FAILED":
                    return None
    except Exception as e:
        print(f"❌ 异常: {e}")
        return None

def clean_json(raw_input):
    text = ""
    if isinstance(raw_input, dict):
        text = raw_input.get("content") or raw_input.get("output") or raw_input.get("result")
        if not text:
            if "header_info" in raw_input or "meta" in raw_input: return raw_input
            text = json.dumps(raw_input)
    else:
        text = str(raw_input)

    if text:
        text = re.sub(r'```json\s*', '', text)
        text = re.sub(r'```\s*', '', text)

    start, end = text.find('{'), text.rfind('}')
    if start != -1 and end != -1:
        try: return json.loads(text[start : end + 1])
        except: pass
        
    # 💡 核心修改：如果解析JSON失败，说明AI返回的是纯文本，那就直接返回这个文本！
    return text if text else None 


# ================= 界面初始化 =================

st.set_page_config(page_title="AI 研报生成器", page_icon="📄", layout="wide")

# 🌟 修改点：初始化 Session State 以保存生成的摘要历史
if "summary_history" not in st.session_state:
    st.session_state.summary_history = []

st.title("📄 AI 智能研报工作台")
st.markdown("上传 PDF -> AI 提取分析 -> 生成报告或摘要")

# --- 侧边栏配置 ---
with st.sidebar:
    st.header("⚙️ 报告配置")
    user_name = st.text_input("用户名称 (User Name)", value="Charlotte")
    report_category = st.selectbox(
        "报告类别 (Category)",
        ("Equity", "Macro", "FX&Commodities","Weekly Fund Flow"),
        index=0
    )
    st.info(f"当前模式: {report_category}\n(Equity 会包含股价评级，其他则隐藏)")


# 🌟 核心修改点：创建三标签页
tab1, tab2, tab3 = st.tabs(["📝 Word 报告生成", "📑 批量中文摘要生成", "💎 投行精髓汇总"])


# ================= TAB 1: 现有的 Word 报告生成逻辑 =================
with tab1:
    st.subheader("生成排版精美的 Word 研报")
    uploaded_pdf = st.file_uploader("上传单份 PDF 研报", type=["pdf"], key="single_word_pdf")
    
    uploaded_image_manual = None
    if report_category == "Weekly Fund Flow":
        st.caption("✅ 资金流模式无需封面图。")
    else:
        uploaded_image_manual = st.file_uploader("上传封面图 (可选)", type=["png", "jpg", "jpeg"])

    generate_btn = st.button("🚀 开始生成 Word 报告", type="primary")

    if generate_btn and uploaded_pdf:
        status_box = st.status("正在处理...", expanded=True)
        try:
            status_box.write("📄 正在读取 PDF 内容...")
            pdf_text = extract_pdf_text(uploaded_pdf)

            if not pdf_text:
                status_box.update(label="❌ PDF 读取失败或为空", state="error")
                st.stop()
                
            if report_category == "Weekly Fund Flow":
                status_box.write("🔍 Step 1: 提取资金流数据...")
                raw_data = call_ai_and_wait_generic(FUND_FLOW_STEP1, pdf_text)
                if not raw_data: 
                    status_box.update(label="❌ Step 1 失败", state="error")
                    st.stop()

                status_box.write("✍️ Step 2: 执行翻译标准...")
                final_json = call_ai_and_wait_generic(FUND_FLOW_STEP2, json.dumps(raw_data))
                
                allowed_keys = ["title", "summary", "body_content", "date", "from", "language"]
                header = final_json.get("header_info", {})
                for key in list(header.keys()):
                    if key.lower() not in allowed_keys:
                        header[key] = ""

                bank_acronym = "GS"
                final_filename = f"WeeklyFlow_{user_name}_{bank_acronym}_{datetime.now().strftime('%Y%m%d')}.docx"
                
                today_str = datetime.now().strftime("%Y/%m/%d")
                if "header_info" in final_json:
                    final_json["header_info"]["date"] = today_str

                status_box.write("💾 正在生成 Word 文档...")
                generator = DocGenerator()
                output_docx_path = f"temp_{final_filename}"
                temp_img_path = None
                generator.create_styled_doc(final_json, output_docx_path, img_path=None, report_category=report_category)
                status_box.update(label="✅ 生成成功！", state="complete", expanded=False)

            else:
                status_box.write("🧠 AI Step 1: 正在提取关键数据...")
                prompt_1 = STEP_1_PROMPT_TEMPLATE.format(category=report_category)
                raw_data = call_ai_and_wait_generic(prompt_1, pdf_text)

                if not raw_data:
                    status_box.update(label="❌ 第一步 AI 分析失败", state="error")
                    st.stop()

                status_box.write("✍️ AI Step 2: 正在进行格式化...")
                prompt_2 = STEP_2_PROMPT_TEMPLATE.format(category=report_category)
                step1_str = json.dumps(raw_data, indent=2, ensure_ascii=False)
                final_json = call_ai_and_wait_generic(prompt_2, step1_str)

                today_str = datetime.now().strftime("%Y/%m/%d")
                if "header_info" in final_json:
                    final_json["header_info"]["date"] = today_str

                original_filename = os.path.splitext(uploaded_pdf.name)[0]
                institution = raw_data.get("meta", {}).get("institution", "Unknown")
                bank_acronym = get_bank_acronym(institution)
                final_filename = f"{report_category}_{user_name}_{bank_acronym}_{original_filename}.docx"
                final_filename = final_filename.replace(" ", "_").replace("/", "-")

                temp_img_path = None
                if uploaded_image_manual:
                    temp_img_path = f"temp_{uploaded_image_manual.name}"
                    with open(temp_img_path, "wb") as f:
                        f.write(uploaded_image_manual.getbuffer())
                    status_box.write(f"🖼️ 已加载封面图: {uploaded_image_manual.name}")

                status_box.write("💾 正在生成 Word 文档...")
                generator = DocGenerator()
                output_docx_path = f"temp_{final_filename}"
                generator.create_styled_doc(final_json, output_docx_path, img_path=temp_img_path, report_category=report_category)
                status_box.update(label="✅ 生成成功！", state="complete", expanded=False)

            # 显示下载按钮
            with open(output_docx_path, "rb") as f:
                st.download_button(
                    label=f"⬇️ 下载报告: {final_filename}",
                    data=f.read(),
                    file_name=final_filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

            # 清理临时文件
            if os.path.exists(output_docx_path): os.remove(output_docx_path)
            if temp_img_path and os.path.exists(temp_img_path): os.remove(temp_img_path)

        except Exception as e:
            status_box.update(label="❌ 发生未知错误", state="error")
            st.error(f"Error details: {e}")

    elif generate_btn and not uploaded_pdf:
        st.warning("请先上传 PDF 文件！")


# ================= TAB 2: 追加的批量中文摘要生成逻辑 =================
with tab2:
    st.subheader("中文摘要批量生成")
    st.info("📌 一次性上传最多 4 份 PDF，系统将为每份生成精炼的中文摘要")
    st.divider()
    
    left_col, right_col = st.columns([4, 6], gap="large")
    
    with left_col:
        st.markdown("#### ⚙️ 全局配置")
        # 提供了模型选择下拉框
        selected_model_tab2 = st.selectbox("🤖 选择 AI 模型", AVAILABLE_MODELS, index=0)
        st.write("") 
        
        st.markdown("#### 📁 上传文件")
        uploaded_pdfs = st.file_uploader(
            "上传PDF研报（可多选，最多4份）",
            type=["pdf"],
            accept_multiple_files=True,
            key="summary_pdfs"
        )
        
        report_configs = []
        generate_summary_btn = False
        
        if uploaded_pdfs:
            if len(uploaded_pdfs) > 4:
                st.error("❌ 最多只能上传 4 份 PDF！")
            else:
                st.success(f"✅ 已选择 {len(uploaded_pdfs)} 份报告")
                st.markdown("#### 🔗 配置链接 ID")
                inner_cols = st.columns(2)
                for idx, pdf_file in enumerate(uploaded_pdfs):
                    with inner_cols[idx % 2]:
                        link_id = st.text_input(
                            f"📎 {pdf_file.name}",
                            placeholder="输入 ID",
                            key=f"link_id_{idx}"
                        )
                        report_configs.append({
                            "pdf": pdf_file,
                            "link_id": link_id
                        })
                
                st.write("")
                generate_summary_btn = st.button("🚀 生成中文摘要", type="primary", key="summary_btn", use_container_width=True)

    with right_col:
        st.markdown("#### ⚙️ 运行状态与结果")
        
        if generate_summary_btn:
            if any(not cfg["link_id"] for cfg in report_configs):
                st.error("❌ 请为所有 PDF 填写链接 ID!")
            else:
                status_msg = st.empty()
                detail_msg = st.empty()
                progress_bar = st.progress(0)
                
                status_msg.info("📄 正在读取所有 PDF 文件...")
                all_pdfs_data = []
                
                for idx, config_item in enumerate(report_configs):
                    pdf_file = config_item["pdf"]
                    link_id = config_item["link_id"]
                    try:
                        pdf_text = extract_pdf_text(pdf_file)
                        if not pdf_text:
                            detail_msg.error(f"❌ 无法读取 {pdf_file.name}")
                            continue
                        all_pdfs_data.append({
                            "name": pdf_file.name,
                            "text": pdf_text,
                            "link_id": link_id
                        })
                    except Exception as e:
                        detail_msg.error(f"❌ 读取 {pdf_file.name} 时出错: {e}")
                
                if not all_pdfs_data:
                    status_msg.error("❌ 无法读取任何 PDF 文件")
                else:
                    status_msg.success(f"✅ 成功读取 {len(all_pdfs_data)} 份 PDF")
                    status_msg.info(f"🧠 正在调用 AI ({selected_model_tab2}) 生成中文摘要...")
                    summary_results = []
                    
                    for idx, pdf_data in enumerate(all_pdfs_data):
                        pdf_name = pdf_data["name"]
                        pdf_text = pdf_data["text"]
                        link_id = pdf_data["link_id"]
                        
                        try:
                            detail_msg.info(f"⏳ 正在处理第 {idx+1}/{len(all_pdfs_data)} 份: {pdf_name}")
                            ai_input = f"Report Filename: {pdf_name}\n\n{pdf_text}\n\n---\nReport Link ID: {link_id}"
                            
                            # 🌟 调用时传入了用户指定的模型
                            summary_data = call_ai_and_wait_generic(
                                getattr(config, "STEP_3_PROMPT", "Summarize this report in Chinese."), 
                                ai_input,
                                model_name=selected_model_tab2
                            )
                            
                            if summary_data:
                                if isinstance(summary_data, dict) and "link_id" not in summary_data:
                                    summary_data["link_id"] = int(link_id)
                                summary_results.append({
                                    "name": pdf_name,
                                    "data": summary_data,
                                    "link_id": link_id 
                                })
                            else:
                                detail_msg.error(f"❌ {pdf_name} AI处理失败")
                        except Exception as e:
                            detail_msg.error(f"❌ 处理 {pdf_name} 时出错: {e}")
                        
                        progress_bar.progress((idx + 1) / len(all_pdfs_data))
                    
                    detail_msg.empty()
                    status_msg.success("✨ 本次任务处理完毕！")
                    
                    if summary_results:
                        summary_results.sort(key=lambda x: int(x.get("link_id", 0)) if str(x.get("link_id", 0)).isdigit() else 0)
                        plain_text_output = []
                        for idx, result_item in enumerate(summary_results):
                            result = result_item["data"]
                            if isinstance(result, dict):
                                title = result.get('title', '未知标题')
                                preview = result.get('preview', '无摘要内容')
                                link = result.get('link', link_id)
                                text_block = f"{title}\n{preview}\n链接：{link}\n"
                                plain_text_output.append(text_block)
                            else:
                                plain_text_output.append(str(result) + "\n")
                        
                        full_text = "\n".join(plain_text_output)
                        
                        st.session_state.summary_history.insert(0, {
                            "time": datetime.now().strftime("%H:%M:%S"),
                            "model": selected_model_tab2,
                            "content": full_text
                        })

        if not st.session_state.summary_history and not generate_summary_btn:
            st.caption("👈 请先在左侧选择模型、上传文件、配置链接 ID，并点击“生成”按钮。")
            
        if st.session_state.summary_history:
            col_title, col_btn = st.columns([8, 2])
            with col_title:
                st.markdown("##### 📚 生成历史 (最新在最前)")
            with col_btn:
                if st.button("🗑️ 清空历史", use_container_width=True):
                    st.session_state.summary_history = []
                    st.rerun() 
            
            for idx, record in enumerate(st.session_state.summary_history):
                st.markdown(f"**🤖 模型: `{record['model']}`** ⏱️ 时间: {record['time']}")
                st.code(record['content'], language="text")
                st.divider()

# ================= TAB 3: 投行精髓汇总生成逻辑 =================
with tab3:
    st.subheader("💎 投行精髓汇总生成 (专业分步版)")
    st.info("📌 架构升级：请分别上传【官方财报】和【各大投行研报】。AI 将先提取绝对准确的财务数据，再逐一深度解析各家投行观点。")
    st.divider()
    
    left_col_t3, right_col_t3 = st.columns([4, 6], gap="large")
    
    with left_col_t3:
        st.markdown("#### ⚙️ 全局配置")
        selected_model_tab3 = st.selectbox("🤖 选择 AI 模型", AVAILABLE_MODELS, index=0, key="model_tab3_new")
        company_name = st.text_input("🏢 目标公司名称", placeholder="例如：腾讯控股（0700.HK）")
        st.write("") 
        
        st.markdown("#### 📁 上传文件 (分步上传)")
        uploaded_official_pdf = st.file_uploader(
            "1️⃣ 上传公司【官方财报】或新闻稿 PDF (单选)",
            type=["pdf"],
            key="wsh_official_pdf"
        )
        
        uploaded_bank_pdfs = st.file_uploader(
            "2️⃣ 上传各大【投行研报】 PDF (可多选)",
            type=["pdf"],
            accept_multiple_files=True,
            key="wsh_bank_pdfs"
        )
        
        st.write("")
        generate_wsh_btn = st.button("🚀 开始分步深度生成", type="primary", key="wsh_btn_new", use_container_width=True)

    with right_col_t3:
        st.markdown("#### ⚙️ 运行状态与结果")
        
        
        if generate_wsh_btn:
            if not uploaded_official_pdf or not uploaded_bank_pdfs:
                st.error("❌ 请确保【官方财报】和【投行研报】都已上传文件！")
            elif not company_name:
                st.warning("⚠️ 建议填写【目标公司名称】，以免 AI 抓取标题失败。")
            else:
                status_msg_t3 = st.empty()
                detail_msg_t3 = st.empty()
                progress_bar = st.progress(0)
                
                # 存放生成结果的变量
                part_performance = "" # 业绩回顾
                part_banks = []       # 投行观点数组
                
                # --- 阶段 1：只提炼业绩回顾 ---
                status_msg_t3.info("📊 阶段 1/3：正在提取【官方财报】硬核数据...")
                official_text = extract_pdf_text(uploaded_official_pdf)
                if not official_text:
                    st.stop()
                
                res_a = call_ai_and_wait_generic(config.PROMPT_A_OFFICIAL, f"目标公司：{company_name}\n\n{official_text}", model_name=selected_model_tab3)
                if res_a: part_performance = str(res_a)
                progress_bar.progress(0.2)
                
                # --- 阶段 2：逐个拆解投行观点 ---
                status_msg_t3.info("🏦 阶段 2/3：正在逐个解析【投行观点】...")
                for idx, bank_pdf in enumerate(uploaded_bank_pdfs):
                    detail_msg_t3.info(f"读取投行报告: {bank_pdf.name}")
                    bank_text = extract_pdf_text(bank_pdf)
                    if bank_text:
                        res_b = call_ai_and_wait_generic(config.PROMPT_B_BANK, f"目标公司：{company_name}\n\n{bank_text}", model_name=selected_model_tab3)
                        if res_b: part_banks.append(str(res_b))
                    
                    progress_bar.progress(0.2 + 0.6 * ((idx + 1) / len(uploaded_bank_pdfs)))

                # --- 阶段 3：画龙点睛（全局摘要生成） ---
                status_msg_t3.info("👑 阶段 3/3：正在站在上帝视角撰写【主标题与核心摘要】...")
                detail_msg_t3.info("综合各大投行观点生成共识中...")
                
                # 把前两步的内容拼起来喂给 Prompt C
                synthesis_context = f"目标公司：{company_name}\n\n【已经提取的业绩回顾】\n{part_performance}\n\n【已经提取的各大投行观点】\n" + "\n\n".join(part_banks)
                
                part_summary = call_ai_and_wait_generic(config.PROMPT_C_SUMMARY, synthesis_context, model_name=selected_model_tab3)
                progress_bar.progress(1.0)
                
                # --- 终极拼装大一统 ---
                detail_msg_t3.empty()
                status_msg_t3.success("✨ 所有步骤处理完毕！正在生成排版文档...")
                
                # 完美的顺序：标题与摘要 -> 业绩回顾 -> 各大投行
                full_report_content = f"{str(part_summary)}\n\n{part_performance}\n\n" + "\n\n".join(part_banks)
                
                # 预览框
                with st.expander("👀 点击预览生成的纯文本内容"):
                    st.markdown(full_report_content)
                
                # ==========================================
                # 下面的 Word 排版和下载代码与之前保持完全一致
                               # --- 极致排版：全文等线、11号、两端对齐、特定标红加粗 ---
                # --- 极致排版：精准空行、大标题居中、全文等线、11号、特定标红 ---
                doc = Document()
                for line in full_report_content.split('\n'):
                    line = line.strip()
                    
                    # 忽略 Markdown 自带的杂乱空行，由我们用代码精准控制空行
                    if not line:
                        continue
                    
                    # 💡 核心逻辑：在“业绩回顾”、“投行标题”、“估值”、“潜在风险”之前，强行插入一个空段落（空一行）
                    if line.startswith('## ') or line.startswith('### ') or line in ["估值", "潜在风险", "**估值**", "**潜在风险**", "估值：", "潜在风险："]:
                        doc.add_paragraph()
                        
                    p = doc.add_paragraph()
                    
                    # 1. 拦截主标题 (#) -> 居中、加粗、后续空一行
                    if line.startswith('# '):
                        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER # 强制居中对齐
                        clean_title = line.replace('# ', '').strip()
                        run = p.add_run(clean_title)
                        run.bold = True
                        run.font.color.rgb = RGBColor(0, 0, 0)
                        doc.add_paragraph() # 主标题写完后，立刻在下方空一行 (留给概括)
                        
                    # 2. 拦截业绩回顾 (##) -> 加粗、两端对齐
                    elif line.startswith('## '):
                        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        clean_title = line.replace('## ', '').strip()
                        run = p.add_run(clean_title)
                        run.bold = True
                        run.font.color.rgb = RGBColor(0, 0, 0)
                        
                    # 3. 拦截投行标题 (###) -> 标红、加粗、两端对齐
                    elif line.startswith('### '):
                        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        clean_title = line.replace('### ', '').strip()
                        run = p.add_run(clean_title)
                        run.bold = True
                        run.font.color.rgb = RGBColor(192, 0, 0) # 保持深红色
                        
                    # 4. 拦截估值与潜在风险标签 -> 加粗、两端对齐
                    elif line in ["估值", "潜在风险", "**估值**", "**潜在风险**", "估值：", "潜在风险："]:
                        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        clean_text = line.replace('**', '').strip()
                        run = p.add_run(clean_text)
                        run.bold = True
                        run.font.color.rgb = RGBColor(0, 0, 0)
                        
                    # 5. 处理正文及列表分论点 -> 两端对齐
                    else:
                        is_list_item = line.startswith('- ') or line.startswith('* ') or line.startswith('·')
                        if is_list_item:
                            # 去掉列表前缀，使用 Word 原生的项目符号
                            clean_line = line.replace('- ', '', 1).replace('* ', '', 1).replace('·', '', 1).strip()
                            p.style = 'List Bullet' 
                            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        else:
                            clean_line = line
                            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        
                        # 处理句子中间可能加粗的重点词汇
                        parts = re.split(r'(\*\*.*?\*\*)', clean_line)
                        for part in parts:
                            if part.startswith('**') and part.endswith('**'):
                                run = p.add_run(part[2:-2])
                                run.bold = True
                                run.font.color.rgb = RGBColor(0, 0, 0)
                            else:
                                run = p.add_run(part)
                                run.font.color.rgb = RGBColor(0, 0, 0)
                
                # 👑 终极格式覆盖：锁定等线和 11 号字（不碰对齐和颜色）
                for paragraph in doc.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'DengXian'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '等线')
                        run.font.size = Pt(11)

                word_buffer = io.BytesIO()
                doc.save(word_buffer)
                word_buffer.seek(0)
                
                st.download_button(
                    label="⬇️ 下载排版好的 Word (.docx) 汇总报告",
                    data=word_buffer,
                    file_name=f"投行精髓_{company_name}_{datetime.now().strftime('%Y%m%d')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    type="primary"
                )
