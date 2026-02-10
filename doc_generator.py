import re
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches 
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

# 指定红色 (RGB: 192, 0, 0)
CUSTOM_RED = RGBColor(192, 0, 0) 

class DocGenerator:
    def create_styled_doc(self, json_data, output_path="Output.docx", img_path=None, report_category=None):
        if not json_data:
            print("❌ 数据为空，无法生成")
            return

        doc = Document()

        # --- 基础字体设置 ---
        style = doc.styles['Normal']
        style.font.name = 'DengXian'
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '等线 (中文正文)') 
        
        # 🔥 根据截图要求：如果是 Weekly Fund Flow，字号设为 14 
        if report_category == "Weekly Fund Flow":
            style.font.size = Pt(14)
        else:
            style.font.size = Pt(11)

        # --- 辅助函数：段落排版 (根据截图优化) ---
        def apply_paragraph_style(paragraph, align_justify=True):
            pf = paragraph.paragraph_format
            pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if align_justify else WD_ALIGN_PARAGRAPH.LEFT
            
            # 🔥 匹配截图设置:
            if report_category == "Weekly Fund Flow":
                pf.space_before = Pt(0)   # 段前: 0 磅
                pf.space_after = Pt(8)    # 段后: 8 磅
                pf.line_spacing = 1.08    # 设置值: 1.08
            else:
                pf.space_before = Pt(12)
                pf.space_after = Pt(0)
                pf.line_spacing = 1.07
                
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE # 多倍行距

        # --- 核心函数：标红重点句 ---
        def add_paragraph_with_highlight(document, text):
            p = document.add_paragraph()
            apply_paragraph_style(p)
            segments = re.split(r'(\*\*.*?\*\*)', str(text))
            for seg in segments:
                if not seg: continue
                if seg.startswith('**') and seg.endswith('**'):
                    clean_text = seg[2:-2]
                    run = p.add_run(clean_text)
                    run.font.color.rgb = CUSTOM_RED
                else:
                    run = p.add_run(seg)
                    run.font.color.rgb = RGBColor(0, 0, 0)

        # --- 1. 顶部信息 (Header) ---
        header = json_data.get("header_info", {})
        fund_flow_mapping = [
            ("Category", ["Category", "category"]),
            ("Date", ["Date", "date"]),
            ("Title", ["Title", "title"]),
            ("Summary", ["Summary", "summary"]),
            ("From", ["From", "from"]),
            ("Tags", ["Tags", "tags"]),
            ("Recommend Expire Time", ["Recommend Expire Time", "expire_time"]),
            ("Language", ["Language", "language"]),
            ("Stock", ["Stock", "stock"]),
            ("Stock Rating", ["Stock Rating", "rating"]),
            ("12m Price Target", ["12m Price Target", "price_target"]),
            ("Related Stock List", ["Related Stock List", "related_stocks"]),
            ("Related Stock Rating", ["Related Stock Rating", "related_rating"])
        ]


        for label, keys in fund_flow_mapping:
            val = ""
            for k in keys:
                if k in header:
                    val = header[k]
                    break
            
            # Weekly Fund Flow 模式强制显示所有标签 
            if report_category == "Weekly Fund Flow":
                p = doc.add_paragraph()
                apply_paragraph_style(p)
                run = p.add_run(f"#{label}# ")
                run.font.bold = False
                if val:
                    run_val = p.add_run(str(val))
                    run_val.font.bold = False
                   # --- 2. 正文 (Content) ---
            
            elif val:
                p = doc.add_paragraph()
                apply_paragraph_style(p, align_justify=True)
                # 标签部分
                run = p.add_run(f"#{label}# ")
                run.font.bold = True
                # 数值部分
                run_val = p.add_run(str(val))
                run_val.font.bold = True # Header部分全部加粗
                   # --- 2. 正文 (Content) ---
        display_name = report_category if report_category else "Market"
        p = doc.add_paragraph()
        apply_paragraph_style(p)
        if report_category == "Weekly Fund Flow":      
            run = p.add_run("#Content#")
            run.font.bold = False
        else :
            run = p.add_run("#Content#")
            run.font.bold = True
          # 2. 创建一个新段落，实现“下一行”的效果
            p_next = doc.add_paragraph()
            apply_paragraph_style(p_next)
            # 写入 Wall Street Highlights 内容
            run_highlight = p_next.add_run(f"Wall Street Highlights-{display_name}")
            run_highlight.font.bold = True
        print(report_category)
        if report_category:
            clean_category = report_category
        else:
            # 尝试从 JSON 获取
            json_cat_raw = header.get("category", "")
            if json_cat_raw:
                # 无论 JSON 里是 "Equity" 还是 "Wall Street Highlights-Equity"，我们都清洗一下
                clean_category = json_cat_raw.replace("Wall Street Highlights-", "").strip()

        # 生成显示的文字
        highlight_title = f"Wall Street Highlights-{clean_category}"

        p = doc.add_paragraph()
        apply_paragraph_style(p)
        
        if report_category == "Weekly Fund Flow":      
            run = p.add_run("#Content#")
            run.font.bold = False
        else:
            # 1. 先写 #Content#
            run = p.add_run("#Content#")
            run.font.bold = True
            
            # 2. 另起一行，写 Wall Street Highlights-Equity (不加 #，加粗)
            p_next = doc.add_paragraph()
            apply_paragraph_style(p_next)
            
            # 这里直接用我们清洗好的名字
            run_highlight = p_next.add_run(highlight_title)
            run_highlight.font.bold = True
            
        body_content = json_data.get("body_content", [])
        body_list = [body_content] if isinstance(body_content, str) else body_content

        for paragraph_text in body_list:
            if paragraph_text.strip():
                add_paragraph_with_highlight(doc, paragraph_text)
                    
        
        if report_category != "Weekly Fund Flow":
            # --- 3. 底部信息 (Footer) - 全红 ---
            footer = json_data.get("footer_info", {})
            if footer:
                footer_items = [
                    ("Stock", footer.get("stock", "")),
                    ("Stock Rating", footer.get("rating", "")),
                    ("12m Price Target", footer.get("price_target", ""))
                ]
                for label, val in footer_items:
                    if val:
                        p = doc.add_paragraph(f"{label}: {val}")
                        apply_paragraph_style(p, align_justify=False)
                        for run in p.runs:
                            run.font.bold = True
                            run.font.color.rgb = CUSTOM_RED # 🔴 底部也用同一个红色

        # --- 3. 插入图片 ---
        if img_path and os.path.exists(img_path):
            img_p = doc.add_paragraph()
            img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER 
            img_p.paragraph_format.space_before = Pt(24) 
            run = img_p.add_run()
            run.add_picture(img_path, width=Inches(6.0))

        doc.save(output_path)






